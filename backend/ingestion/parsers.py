"""Document parsers for supported upload types."""

from __future__ import annotations

import asyncio
from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from fastapi import HTTPException, status
from pypdf import PdfReader


@dataclass(frozen=True)
class ParsedDocument:
    """Parsed document text plus coarse page count."""

    text: str
    page_count: int


def _parse_pdf(path: Path, max_pages: int) -> ParsedDocument:
    reader = PdfReader(str(path))
    page_count = len(reader.pages)
    if page_count > max_pages:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"PDF has {page_count} pages; the per-document limit is {max_pages} pages.",
        )
    pages: list[str] = []
    for page_index in range(min(page_count, max_pages)):
        page_number = page_index + 1
        page = reader.pages[page_index]
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"\n\n[Page {page_number}]\n{text}")
    return ParsedDocument(text="\n".join(pages), page_count=page_count)


def _parse_docx(path: Path) -> ParsedDocument:
    doc = DocxDocument(str(path))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return ParsedDocument(text="\n".join(parts), page_count=1)


def _parse_txt(path: Path) -> ParsedDocument:
    return ParsedDocument(text=path.read_text(encoding="utf-8", errors="ignore"), page_count=1)


async def parse_document(path: Path, max_pages: int) -> ParsedDocument:
    """Parse text from a supported document."""

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return await asyncio.to_thread(_parse_pdf, path, max_pages)
    if suffix == ".docx":
        return await asyncio.to_thread(_parse_docx, path)
    if suffix == ".txt":
        return await asyncio.to_thread(_parse_txt, path)
    raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=f"Unsupported file type: {suffix}")
